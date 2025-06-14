import streamlit as st
import re
from docx import Document
import pandas as pd
import unicodedata
import os
import io
import zipfile

def normalize(text):
    return unicodedata.normalize("NFKC", text.replace('\xa0', ' ').replace('\u200b', '')).strip()

def extract_tags_from_docx(docx_file) -> set:
    pattern = re.compile(r"\{\{\s*(.*?)\s*\}\}")
    jinja_blocks = re.compile(r"{%\s*(if|endif|for|endfor)[^%]*%}")
    doc = Document(docx_file)
    tags = set()
    jinja_found = False

    def check_text(text):
        nonlocal jinja_found
        for match in pattern.findall(text):
            tags.add(normalize(match))
        if jinja_blocks.search(text):
            jinja_found = True

    for p in doc.paragraphs:
        check_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                check_text(cell.text)
    for section in doc.sections:
        for part in [section.header, section.footer]:
            for p in part.paragraphs:
                check_text(p.text)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        check_text(cell.text)

    return tags, jinja_found

def replace_placeholders_in_doc(template, mapping, row):
    """
    Remplace les placeholders {{Tag}} même s’ils sont répartis sur plusieurs runs,
    en conservant le style du premier run de chaque occurrence.
    """
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        if not runs:
            return

        # Reconstruction des textes de runs
        text_runs = [run.text for run in runs]
        full_text = ''.join(text_runs)

        for tag, col in mapping.items():
            if not col or col == "(laisser inchangée)" or col not in row.index:
                continue
            value = str(row[col])
            regex = re.compile(r"\{\{\s*" + re.escape(tag) + r"\s*\}\}")

            # Remplacer toutes les occurrences de la balise
            while True:
                match = regex.search(full_text)
                if not match:
                    break
                start, end = match.start(), match.end()

                # Identifier les runs concernés
                run_positions = []
                pos = 0
                for i, txt in enumerate(text_runs):
                    if pos + len(txt) > start and pos < end:
                        run_positions.append(i)
                    pos += len(txt)

                if not run_positions:
                    break

                first_run = runs[run_positions[0]]
                # Vider le texte des runs concernés
                for idx in run_positions:
                    runs[idx].text = ""
                # Insérer la valeur dans un nouveau run
                new_run = paragraph.add_run(value)
                new_run.bold      = first_run.bold
                new_run.italic    = first_run.italic
                new_run.underline = first_run.underline
                new_run.font.name = first_run.font.name
                new_run.font.size = first_run.font.size

                # Recalculer text_runs et full_text pour les prochaines itérations
                text_runs = [r.text for r in runs]
                full_text = ''.join(text_runs)

    def process(container):
        for p in container.paragraphs:
            replace_in_paragraph(p)
        for table in container.tables:
            for r in table.rows:
                for cell in r.cells:
                    process(cell)

    process(template)
    for section in template.sections:
        process(section.header)
        process(section.footer)

def main():
    # Configuration de la page pour Bêta V3.5
    st.set_page_config(
        page_title="🛠️ Bêta Juridique – Générateur V3.5",
        page_icon="⚖️"
    )
    st.title("🛠️ Bêta Juridique – Assistant de génération V3.5")

    # Initialisation de l'état de mapping
    if "mapping_done" not in st.session_state:
        st.session_state["mapping_done"] = False

    st.markdown("""
    Ce service vous permet de générer automatiquement des documents juridiques  
    à partir d’un modèle Word (.docx) et d’un fichier Excel contenant les informations clients.
    """)

    with st.expander("🔐 Politique de confidentialité"):
        st.markdown("""
        Les fichiers que vous chargez ne sont jamais stockés. Ils sont traités  
        uniquement pendant votre session et supprimés ensuite automatiquement.  
        ✅ Conforme au RGPD.
        """)

    # Upload des fichiers
    word_file  = st.file_uploader("📄 Télécharger votre modèle Word (.docx)", type="docx")
    excel_file = st.file_uploader("📊 Importer votre tableau Excel (.xls/.xlsx)", type=["xls", "xlsx"])

    mapping = {}
    tags = set()
    df = None
    jinja_found = False

    # Extraction des balises
    if word_file:
        tags, jinja_found = extract_tags_from_docx(word_file)
    if jinja_found:
        st.warning("⚠️ Le modèle contient des blocs conditionnels Jinja non traités.")

    # Aperçu des données importées
    if word_file or excel_file:
        with st.expander("📂 Aperçu des données importées"):
            if tags:
                st.markdown("### Champs personnalisables détectés")
                for tag in sorted(tags):
                    st.write(f"- **{{{{{tag}}}}}**")
            elif word_file:
                st.info("Aucune balise {{…}} détectée.")
            if excel_file:
                df = pd.read_excel(excel_file)
                df.columns = df.columns.str.strip()
                st.markdown("### Colonnes disponibles")
                st.write(list(df.columns))

    # Mapping balises → colonnes (strict par défaut, ou tolérant si activé)
    if word_file and excel_file:
        if df is None:
            df = pd.read_excel(excel_file)
            df.columns = df.columns.str.strip()

        st.markdown("### Associer chaque champ du modèle aux données Excel")

        # Toggle strict vs tolérant
        tol = st.checkbox(
            "Activer le mapping tolérant (ignorer la casse et les underscores)",
            value=False
        )

        cols = ["(laisser inchangée)"] + list(df.columns)
        normalized = [c.lower().replace("_", "") for c in df.columns]

        for tag in sorted(tags):
            if tol:
                tag_norm = tag.lower().replace("_", "")
                default = normalized.index(tag_norm) + 1 if tag_norm in normalized else 0
            else:
                default = cols.index(tag) if tag in df.columns else 0

            mapping[tag] = st.selectbox(f"Champ modèle : {{{{{tag}}}}}", cols, index=default)

        if st.button("🔗 Enregistrer les correspondances"):
            st.session_state["mapping_done"] = True
            st.success("🔄 Correspondances enregistrées avec succès.")

    # Génération et téléchargement
    if word_file and excel_file and st.session_state["mapping_done"]:
        if st.button("📂 Générer les documents personnalisés"):
            # Lecture et nettoyage
            df = pd.read_excel(excel_file)
            df.columns = df.columns.str.strip()

            # Préparation du nom du modèle
            raw        = os.path.splitext(word_file.name)[0]
            clean_name = raw.replace("_", " ")
            parts      = clean_name.split(" ", 1)
            prefix     = parts[0]
            rest       = parts[1] if len(parts) > 1 else ""

            # Extraction major/minor
            try:
                major, minor = prefix.split(".")
                minor = int(minor)
            except ValueError:
                major, minor = prefix, 0

            # Construction du ZIP
            zip_io = io.BytesIO()
            with zipfile.ZipFile(zip_io, "w") as zf:
                for idx, row in df.iterrows():
                    template = Document(word_file)
                    replace_placeholders_in_doc(template, mapping, row)
                    seq      = minor + idx + 1
                    new_pref = f"{major}.{seq}"
                    key      = next((col for tag, col in mapping.items() 
                                     if tag.lower() == "name" and col in row.index), None)
                    person   = str(row[key]).strip() if key else "inconnu"
                    fname    = f"{new_pref} {rest} - {person}.docx"
                    out = io.BytesIO()
                    template.save(out)
                    zf.writestr(fname, out.getvalue())
            zip_io.seek(0)

            # Stockage en session
            st.session_state["zip_data"]     = zip_io.getvalue()
            st.session_state["zip_filename"] = f"{clean_name}.zip"

        # Bouton de téléchargement visible après génération
        if st.session_state.get("zip_data"):
            st.download_button(
                "📥 Télécharger l’ensemble des documents (ZIP)",
                data=st.session_state["zip_data"],
                file_name=st.session_state["zip_filename"],
                mime="application/zip",
                key="download-zip"
            )

if __name__ == "__main__":
    main()
